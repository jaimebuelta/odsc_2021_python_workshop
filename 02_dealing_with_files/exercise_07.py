import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
import click
import matplotlib.pyplot as plt
from collections import defaultdict


@click.command()
@click.argument('file_in', type=click.File('r'))
@click.argument('file_out', type=click.Path())
def action(file_in, file_out):

    seasons = defaultdict(int)

    for row in csv.DictReader(file_in):
        season = row['Season']
        seasons[season] += 1

    # sort and write the result
    data = list(seasons.items())
    pos = list(range(len(data)))
    values = [value for label, value in data]
    labels = [label for label, value in data]
    plt.bar(pos, values)
    plt.xticks(pos, labels)

    figure = plt. gcf()
    figure. set_size_inches(10, 6)
    GRAPHIC = 'graphic.png'
    plt.savefig(GRAPHIC)

    document = docx.Document()
    document.add_heading('The Simpsons seasons', 0)

    paragraph = document.add_paragraph()
    paragraph.add_run('This is an automatic report generating a text '
                      'and graph showing the number of episodes on '
                      'each season of "The Simpsons"')

    # Add the image
    image = document.add_picture(GRAPHIC)
    image.width = Cm(15)
    image.height = Cm(9)
    # This image is added as a paragraph, align it
    paragraph = document.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_break()
    paragraph.add_run('Graph')

    # Add some text afterwards
    paragraph = document.add_paragraph()
    paragraph.add_run('Italic Footer').italic = True
    paragraph.add_run(', and ')
    paragraph.add_run('Bold Footer').bold = True

    document.save(file_out)


if __name__ == '__main__':
    action()
